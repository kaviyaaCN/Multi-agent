"""
agents/documentation_agent.py
==============================
Documentation Generator Agent
==============================
Produces a complete IEEE-style academic project report in structured sections:
Abstract, Introduction, Objectives, System Architecture, Methodology,
Algorithm/Flow, Implementation, Results, Advantages, Limitations,
Conclusion, Future Scope, and References.
"""

from __future__ import annotations

import time
from dataclasses import dataclass, field
from typing import Optional

from utils.llm_client import get_llm_client
from utils.helpers import clean_llm_output
from utils.logger import logger


@dataclass
class ProjectDocumentation:
    """Container for all 13 sections of the generated academic report."""
    topic: str
    domain: str = ""
    abstract: str = ""
    introduction: str = ""
    objectives: str = ""
    system_architecture: str = ""
    methodology: str = ""
    algorithm: str = ""
    implementation: str = ""
    results_and_discussion: str = ""
    advantages: str = ""
    limitations: str = ""
    conclusion: str = ""
    future_scope: str = ""
    references: list[str] = field(default_factory=list)

    def to_full_text(self) -> str:
        """Return the complete document as a single formatted markdown string."""
        return f"""# {self.topic}
### MCA Project Report | Domain: {self.domain}

---

## Abstract
{self.abstract}

---

## 1. Introduction
{self.introduction}

---

## 2. System Architecture
{self.system_architecture}

---

## 3. Methodology
{self.methodology}

---

## 4. Implementation
{self.implementation}

---

## 5. Conclusion
{self.conclusion}
"""

    def to_docx(self) -> bytes:
        """Export the document as a Microsoft Word (.docx) binary stream."""
        import io
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        docx = Document()
        
        # Title
        title = docx.add_heading(self.topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = docx.add_paragraph(f"MCA Project Report | Domain: {self.domain}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Helper to convert basic markdown into word
        def _add_section(d, title_text, content):
            d.add_heading(title_text, level=1)
            # Extremely basic markdown to paragraph converter
            for paragraph in content.split('\n\n'):
                p = paragraph.strip()
                if not p: continue
                
                if p.startswith('## '):
                    d.add_heading(p[3:], level=2)
                elif p.startswith('### '):
                    d.add_heading(p[4:], level=3)
                elif p.startswith('- ') or p.startswith('* '):
                    for item in p.split('\n'):
                        if item.startswith('- ') or item.startswith('* '):
                            d.add_paragraph(item[2:], style='List Bullet')
                        else:
                            d.add_paragraph(item)
                else:
                    d.add_paragraph(p)
                    
        _add_section(docx, "Abstract", self.abstract)
        _add_section(docx, "1. Introduction", self.introduction)
        _add_section(docx, "2. System Architecture", self.system_architecture)
        _add_section(docx, "3. Methodology", self.methodology)
        _add_section(docx, "4. Implementation", self.implementation)
        _add_section(docx, "5. Conclusion", self.conclusion)
        
        byte_stream = io.BytesIO()
        docx.save(byte_stream)
        return byte_stream.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# System prompt used for all section generation
# ─────────────────────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """You are an expert academic writer helping MCA (Master of Computer Applications) 
final-year students write their project reports. 

Your writing style must be:
- Clear and easy to understand (avoid unnecessary jargon)
- Structured with proper headings and bullet points where asked
- Formal but readable — suitable for both academic submission and viva explanation
- Complete and detailed — never use placeholder text like "[insert here]"
- Always use real examples, realistic values, and concrete explanations

Write exactly what is asked for each section. Do not add extra sections or meta-commentary."""


# ─────────────────────────────────────────────────────────────────────────────
# Section-specific prompt templates
# ─────────────────────────────────────────────────────────────────────────────

_SECTION_TEMPLATES: dict[str, str] = {

    # ── Abstract ─────────────────────────────────────────────────────────────
    "abstract": """Write a short, clear Abstract (100–120 words) for an MCA project titled:
"{topic}"
Domain: {domain} | Difficulty: {difficulty}

Relevant context: {context}

The abstract must cover these 4 points in simple English:
1. What problem does this project solve?
2. What approach / method was used?
3. What tools or techniques were used?
4. What result or outcome was achieved?

Use plain language. Avoid heavy research jargon.
Write ONLY the abstract text — no heading, no labels.""",

    # ── Introduction ─────────────────────────────────────────────────────────
    "introduction": """Write a detailed Introduction section (200–250 words) for an MCA project titled:
"{topic}"
Domain: {domain}

Relevant context: {context}

Structure the introduction with these sub-sections using bold headings:
**What is {domain}?** — Explain the domain in simple terms with a real-world example.
**Problem Statement** — What specific problem does this project address?
**Motivation** — Why is this problem important? Who benefits?
**Proposed Solution** — Brief overview of how the project solves the problem.
**Scope** — What does the project cover and what is out of scope?

Use simple language suitable for MCA level. Include a relatable example.
Write ONLY the introduction text.""",

    # ── Objectives ───────────────────────────────────────────────────────────
    "objectives": """Write the Objectives section for an MCA project titled:
"{topic}"
Domain: {domain}

Write exactly 8 clear objective bullet points.
Format each using: "• To [action verb] [what] [why/result]"

Examples of good objectives:
• To design and implement a [specific component] for [purpose]
• To evaluate the performance of [method] using [metric]
• To compare [approach A] with [approach B] in terms of [criteria]

Make objectives specific, measurable, and relevant to this project.
Write ONLY the bullet-point list — no heading, no intro sentence.""",

    # ── System Architecture ───────────────────────────────────────────────────
    "system_architecture": """Write a detailed System Architecture section (200–250 words) for an MCA project titled:
"{topic}"
Technologies: {technologies}

Include:

**Overview Diagram Description**
Draw a simple text-based ASCII diagram showing the main system components and how data flows between them. Use arrows (→) to show flow. At minimum include: Input → Processing → Output layers.

**Component Table**
Create a table with 3 columns: Component | Description | Technology Used

**Data Flow Explanation**
Explain in 3–5 sentences how data moves through the system from user input to final output.

Use concrete component names based on the actual project technologies.
Write ONLY the architecture content.""",

    # ── Methodology ──────────────────────────────────────────────────────────
    "methodology": """Write a clear Methodology section (250–300 words) for an MCA project titled:
"{topic}"
Technologies: {technologies}
Domain: {domain}

Context: {context}

Structure it as follows:

**Step-by-Step Process**
Explain the methodology as numbered steps (Step 1, Step 2, etc.).
Each step should have:
- A bold heading
- 2–3 sentences explaining what happens in that step
- Mention of the specific tool/technique used

Cover these phases: Data Collection/Input → Preprocessing → Core Processing → Output Generation → Evaluation

**Design Decisions**
Explain in 2–3 sentences WHY specific technologies or approaches were chosen over alternatives.

Keep explanations beginner-friendly. A first-year student should understand each step.
Write ONLY the methodology text.""",

    # ── Algorithm / Flow ──────────────────────────────────────────────────────
    "algorithm": """Write an Algorithm / Flow section for an MCA project titled:
"{topic}"
Domain: {domain}

Include TWO parts:

**Part 1: Pseudocode**
Write the main algorithm as clear pseudocode using plain English.
Format:
FUNCTION name(inputs):
    Step 1: ...
    Step 2: IF condition THEN ...
    Step 3: FOR each item DO ...
    RETURN result
END FUNCTION

Write at least 2 functions — one for the main logic, one for a key sub-process.

**Part 2: Flow Steps**
Below the pseudocode, write a numbered list of 4–6 flow steps that describe how the system works from start to finish, like:
1. User provides input...
2. System preprocesses...
3. Algorithm checks...
...
N. Result is displayed.

Make the pseudocode readable by someone who knows basic programming but not this specific algorithm.
Write ONLY the algorithm content.""",

    # ── Implementation ────────────────────────────────────────────────────────
    "implementation": """Write a detailed Implementation section (200–300 words) for an MCA project titled:
"{topic}"
Technologies: {technologies}
Domain: {domain}

Include:

**Tools and Technologies Table**
Create a table: Tool/Library | Version | Purpose | Why Chosen

**Project File Structure**
Show the folder and file structure as a tree diagram using ASCII art:
project_name/
├── main.py       ← description
├── module.py     ← description
└── ...

**Key Code Snippet**
Write ONE complete, working Python code snippet (10–15 lines) for the most important part of this project.
Include:
- Module docstring
- A class or function with clear docstrings
- Inline comments explaining logic
- Realistic variable names

**How to Run**
Write 3–4 bash commands showing how to install and run the project.

Write ONLY the implementation content.""",

    # ── Results ───────────────────────────────────────────────────────────────
    "results_and_discussion": """Write a Results and Discussion section (200–250 words) for an MCA project titled:
"{topic}"
Domain: {domain}

Include:

**Performance Metrics Table**
Create a table showing realistic evaluation results with at least 4 rows.
Use appropriate metrics for this domain (accuracy, F1, RMSE, precision, recall, etc.).
Use realistic but favourable values (e.g., accuracy around 88–95%).

**Key Findings**
Write 4–5 bullet points summarising what the results show.

**Sample Output**
Describe or show 1 sample input/output example with actual values.

**Discussion**  
Write 2–3 sentences comparing the results to what was expected and what they mean practically.

Use specific numbers throughout. Do not use vague statements like "the results were good".
Write ONLY the results content.""",

    # ── Advantages ────────────────────────────────────────────────────────────
    "advantages": """Write the Advantages section for an MCA project titled:
"{topic}"
Domain: {domain}

Write exactly 8 advantages as bullet points.
Each advantage should:
- Start with a strong emoji (✅ or 🟢) and a bold keyword
- Have 1–2 sentences explaining WHY it is an advantage for this specific project

Example format:
✅ **Fast Processing** — The system processes each request in under 2 seconds, making it suitable for real-time use cases like live recommendations.

Make advantages specific to this project — not generic statements.
Write ONLY the bullet-point advantages.""",

    # ── Limitations ───────────────────────────────────────────────────────────
    "limitations": """Write the Limitations section for an MCA project titled:
"{topic}"
Domain: {domain}

Write exactly 6 limitations as bullet points.
Each limitation should:
- Start with ⚠️ or 🔴 and a bold keyword
- Explain WHAT the limitation is and WHY it exists (1–2 sentences)
- Be honest — real limitations, not trivial ones

Example format:
⚠️ **Limited Dataset** — The model was trained on only 5,000 samples, which may reduce accuracy when applied to data from different sources or regions.

Write ONLY the bullet-point limitations.""",

    # ── Conclusion ────────────────────────────────────────────────────────────
    "conclusion": """Write a clear, concise Conclusion section (220–280 words) for an MCA project titled:
"{topic}"
Domain: {domain}

Structure with these parts:
**Summary** — What was built and what it achieves (2–3 sentences).
**Key Contributions** — 3 numbered contributions that are specific to this project.
**Learning Outcomes** — What skills and knowledge were gained (2 sentences).
**Closing Statement** — One strong closing sentence about the significance of this project.

Keep language simple, positive, and confident.
Write ONLY the conclusion text.""",

    # ── Future Scope ─────────────────────────────────────────────────────────
    "future_scope": """Write a Future Scope section for an MCA project titled:
"{topic}"
Domain: {domain}

Create a table with 3 columns:
| Enhancement | Description | Expected Benefit |

Write exactly 4 rows covering realistic future improvements such as:
- Scaling to handle more data or users
- Adding machine learning / AI capabilities
- Mobile or web app deployment
- Integration with real-world APIs
- Performance optimization
- New features or use cases

Make each enhancement specific and achievable (not vague like "improve the system").
Below the table, write 2–3 sentences describing the overall direction for future development.
Write ONLY the future scope content.""",

    # ── References ────────────────────────────────────────────────────────────
    "references": """Generate 5 realistic, IEEE-format academic references for an MCA project titled:
"{topic}" in the domain of {domain}.

Include a mix of:
- 2 journal papers
- 1 conference paper
- 1 textbook
- 1 online resource (official documentation or reputable website)

IEEE format for journal paper:
A. Author, "Title of Paper," Journal Name, vol. X, no. Y, pp. XXX-XXX, YEAR.

IEEE format for book:
A. Author, Title of Book, Xth ed. City: Publisher, YEAR.

IEEE format for online resource:
A. Author, "Title," Website Name, YEAR. [Online]. Available: https://url

Use realistic author names, titles, and publication details relevant to this domain.
Return ONLY the numbered reference list, one per line, starting with [1].""",
}

# Delay between sections to respect free-tier rate limits
# Delay (seconds) between consecutive section LLM calls.
# Gemini free tier: 20 RPM → 1 request every 3s theoretically,
# but burst limits apply. 20s spacing keeps us well under the limit.
_INTER_SECTION_DELAY = 20

# Section display order
_SECTIONS = [
    "abstract",
    "introduction",
    "system_architecture",
    "methodology",
    "implementation",
    "conclusion",
]


class DocumentationAgent:
    """
    Agent that generates a complete 13-section IEEE / MCA-style project report.

    Each section is generated with a targeted, structured prompt ensuring
    consistent formatting, appropriate depth, and viva-readiness.
    """

    def __init__(self, embedding_engine=None) -> None:
        self._llm = get_llm_client()
        self._engine = embedding_engine
        logger.info(f"DocumentationAgent initialised (RAG={embedding_engine is not None})")

    def generate(
        self,
        topic: str,
        domain: str,
        difficulty: str = "Intermediate",
        technologies: Optional[list[str]] = None,
    ) -> ProjectDocumentation:
        """
        Generate a full 13-section academic project report.

        Args:
            topic: Project title.
            domain: Technical domain.
            difficulty: Beginner / Intermediate / Advanced.
            technologies: List of tools and frameworks used.

        Returns:
            ProjectDocumentation instance with all sections populated.
        """
        technologies = technologies or ["Python", "TensorFlow", "Streamlit"]
        tech_str = ", ".join(technologies)

        logger.info(f"DocumentationAgent.generate: topic={topic}")

        context = self._get_context(f"{topic} {domain} methodology results")

        doc = ProjectDocumentation(topic=topic, domain=domain)

        section_args = {
            "topic": topic,
            "domain": domain,
            "difficulty": difficulty,
            "technologies": tech_str,
            "context": context,
        }

        for i, section in enumerate(_SECTIONS):
            logger.info(f"Generating section: {section} ({i+1}/{len(_SECTIONS)})")
            try:
                prompt = _SECTION_TEMPLATES[section].format(**section_args)
                raw = self._llm.generate(
                    prompt=prompt,
                    system_instruction=_SYSTEM_PROMPT,
                    temperature=0.7,
                    max_tokens=800,
                )
                content = clean_llm_output(raw)

                if section == "references":
                    refs = [line.strip() for line in content.splitlines() if line.strip()]
                    doc.references = refs
                else:
                    setattr(doc, section, content)

                logger.info(f"Section '{section}' complete ({len(content)} chars)")

            except Exception as exc:
                logger.error(f"Failed to generate section '{section}': {exc}")
                if section != "references":
                    setattr(doc, section, f"[Section could not be generated: {exc}]")

            # Rate-limit pause between sections
            if i < len(_SECTIONS) - 1:
                time.sleep(_INTER_SECTION_DELAY)

        logger.info(f"Full report complete: {topic}")
        return doc

    def _get_context(self, query: str) -> str:
        """Retrieve RAG context or return a sensible fallback."""
        if self._engine:
            try:
                return self._engine.get_context(query, top_k=4, max_chars=1500)
            except Exception as exc:
                logger.warning(f"RAG context retrieval failed: {exc}")
        return (
            "Use general best practices, state-of-the-art methods, "
            "and realistic examples relevant to this project domain."
        )
